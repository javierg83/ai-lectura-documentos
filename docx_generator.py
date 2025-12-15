from docx import Document
from docx.shared import Inches
import json


def build_docx_from_json(json_path: str, output_path: str):
    """
    Genera un documento Word (.docx) concatenando el texto y tablas extraídos.
    Aplica márgenes estándar y mantiene saltos de sección al bordear contenido extenso.
    """
    # Cargar datos JSON
    with open(json_path, 'r', encoding='utf-8') as f:
        pages = json.load(f)

    doc = Document()
    # Márgenes estándar (1" = 2.54 cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for page in pages:
        doc.add_paragraph(f"Página {page['pagina']}")
        for elem in page.get('elementos', []):
            if elem.get('tipo') == 'texto':
                # Agregar párrafo; Word manejará el salto de línea dentro de márgenes
                doc.add_paragraph(elem.get('contenido', ''))
            elif elem.get('tipo') == 'tabla':
                rows = elem.get('contenido', [])
                if rows:
                    # Crear tabla con el número de filas y columnas
                    cols = len(rows[0])
                    table = doc.add_table(rows=1, cols=cols)
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(rows[0]):
                        hdr_cells[i].text = str(header)
                    for row_data in rows[1:]:
                        row_cells = table.add_row().cells
                        for i, cell_text in enumerate(row_data):
                            row_cells[i].text = str(cell_text)
        # Salto de página al final de cada página de datos
        doc.add_page_break()

    doc.save(output_path)